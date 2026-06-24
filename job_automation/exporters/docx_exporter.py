from __future__ import annotations

from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config import EXPORT_DIR
from database.models import Job
from exporters.job_presenter import build_job_records


def export_jobs_to_docx(jobs: list[Job], export_dir: Path = EXPORT_DIR) -> Path:
    export_dir.mkdir(parents=True, exist_ok=True)
    path = export_dir / f"jobs_report_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
    records = build_job_records(jobs)

    export_records_to_docx(records, path)
    return path


def export_records_to_docx(records: list[dict[str, str | int]], path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)
    _add_title(document, records)
    _add_summary_table(document, records)
    _add_job_cards(document, records)
    document.save(path)
    return path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)


def _add_title(document: Document, records: list[dict[str, str | int]]) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Job Automation Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_at = records[0]["search_run_at"] if records else date.today().isoformat()
    subtitle_run = subtitle.add_run(f"{len(records)} exported jobs - search run: {run_at}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(75, 85, 99)


def _add_summary_table(document: Document, records: list[dict[str, str | int]]) -> None:
    urgent = sum(1 for record in records if record["priority"] == "urgent")
    high = sum(1 for record in records if record["priority"] == "high")
    avg_score = round(sum(int(record["score"]) for record in records) / len(records), 1) if records else 0
    salary_ok = sum(1 for record in records if record["salary_target_met"] == "yes")
    hours_ok = sum(1 for record in records if record["hours_target_met"] == "yes")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Jobs", "Average Score", "Urgent/High", "Salary OK", "Hours OK"]
    values = [len(records), avg_score, f"{urgent}/{high}", salary_ok, hours_ok]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = f"{header}\n{values[index]}"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    document.add_paragraph()


def _add_job_cards(document: Document, records: list[dict[str, str | int]]) -> None:
    if not records:
        document.add_paragraph("No jobs were exported for the current filters.")
        return

    for index, record in enumerate(records, start=1):
        heading = document.add_paragraph()
        heading.style = document.styles["Heading 2"]
        heading_run = heading.add_run(f"{index}. {record['title']}")
        heading_run.bold = True

        meta = document.add_paragraph()
        meta.add_run(f"{record['company']} - {record['location']} - {record['remote_type']} - Score {record['score']}").bold = True

        detail_table = document.add_table(rows=0, cols=2)
        detail_table.style = "Light Shading Accent 1"
        _add_detail_row(detail_table, "Search run", str(record["search_run_at"]))
        _add_detail_row(detail_table, "Status", str(record.get("status", "new")))
        _add_detail_row(detail_table, "Company type", str(record.get("company_fit", "unknown")))
        _add_detail_row(detail_table, "Company size", str(record.get("company_size", "unknown")))
        _add_detail_row(detail_table, "City / Country", f"{record.get('city', 'Unknown')} / {record.get('country', 'Unknown')}")
        _add_detail_row(detail_table, "Link", str(record["url"]))
        _add_detail_row(detail_table, "Compensation", f"{record['salary']} ({record['salary_hour_notes']})")
        _add_detail_row(detail_table, "Skills", str(record["skills"]))
        if record.get("next_step"):
            _add_detail_row(detail_table, "Next step", str(record["next_step"]))
        if record.get("review_notes"):
            _add_detail_row(detail_table, "Review notes", str(record["review_notes"]))
        _add_detail_row(detail_table, "Why it matches", str(record["reason"]))
        _add_detail_row(detail_table, "Description", str(record["description"]))

        if index != len(records):
            document.add_paragraph()


def _add_detail_row(table, label: str, value: str) -> None:
    row = table.add_row()
    label_cell, value_cell = row.cells
    label_cell.text = label
    value_cell.text = value or "Unknown"
    for paragraph in label_cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
