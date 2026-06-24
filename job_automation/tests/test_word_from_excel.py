from pathlib import Path

from database.models import Job
from docx import Document
from export_word_from_excel import read_records_from_excel
from exporters.docx_exporter import export_records_to_docx
from exporters.xlsx_exporter import export_jobs_to_xlsx


def test_read_records_from_excel_filters_dismissed(tmp_path: Path):
    jobs = [
        Job(
            job_title="Junior AI Automation Specialist",
            company_name="Keep Co",
            location="Remote Germany",
            remote_type="remote",
            job_url="https://example.com/keep",
            relevance_score=90,
            priority_level="urgent",
            reason_for_score="Great fit.",
        ),
        Job(
            job_title="AI Workflow Specialist",
            company_name="Drop Co",
            location="Remote Europe",
            remote_type="remote",
            job_url="https://example.com/drop",
            relevance_score=70,
            priority_level="high",
            reason_for_score="Good fit.",
        ),
    ]
    excel_path = export_jobs_to_xlsx(jobs, tmp_path)

    from openpyxl import load_workbook

    workbook = load_workbook(excel_path)
    sheet = workbook["Job Details"]
    sheet.cell(row=3, column=sheet.max_column, value="yes")
    workbook.save(excel_path)

    active_records = read_records_from_excel(excel_path, include_dismissed=False)
    all_records = read_records_from_excel(excel_path, include_dismissed=True)

    assert len(active_records) == 1
    assert len(all_records) == 2
    assert active_records[0]["company"] == "Keep Co"

    docx_path = export_records_to_docx(active_records, tmp_path / "report.docx")
    assert docx_path.exists()
    document_text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)
    table_text = "\n".join(cell.text for table in Document(docx_path).tables for row in table.rows for cell in row.cells)
    assert "Drop Co" not in document_text + table_text
    assert "Dismissed" not in document_text + table_text
